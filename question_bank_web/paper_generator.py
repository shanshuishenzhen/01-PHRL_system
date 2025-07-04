#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
组题功能核心逻辑
支持按规则自动组题和手动组题
"""

import random
from typing import List, Dict, Optional, Tuple
from sqlalchemy.orm import Session
from models import Question, Paper, PaperQuestion, QuestionGroup, QuestionBank
import datetime
from docx import Document
from docx.shared import Pt
from io import BytesIO

class PaperGenerator:
    """试卷生成器"""
    
    def __init__(self, db_session: Session):
        self.db_session = db_session
        
    def generate_paper_by_rules(self, 
                               paper_name: str,
                               paper_description: str = "",
                               total_score: float = 100.0,
                               duration: int = 120,
                               difficulty_level: str = "中等",
                               rules: Optional[List[Dict]] = None) -> Paper:
        """
        根据规则自动生成试卷
        
        Args:
            paper_name: 试卷名称
            paper_description: 试卷描述
            total_score: 试卷总分
            duration: 考试时长（分钟）
            difficulty_level: 试卷难度等级
            rules: 组题规则列表，格式如下：
                [
                    {
                        "question_type": "B",  # 题型代码
                        "difficulty": "3",     # 难度代码
                        "count": 10,           # 题目数量
                        "score_per_question": 5.0,  # 每题分值
                        "section_name": "单选题"    # 章节名称
                    },
                    ...
                ]
        
        Returns:
            Paper: 生成的试卷对象
        """
        if rules is None:
            rules = self._get_default_rules()
        
        # 创建试卷
        paper = Paper(
            name=paper_name,
            description=paper_description,
            total_score=total_score,
            duration=duration,
            difficulty_level=difficulty_level
        )
        
        self.db_session.add(paper)
        self.db_session.flush()  # 获取paper.id
        
        # 按规则选择题目
        question_order = 1
        for rule in rules:
            questions = self._select_questions_by_rule(rule)
            
            if len(questions) < rule.get('count', 1):
                raise ValueError(f"题型 {rule.get('question_type')} 难度 {rule.get('difficulty')} 的题目数量不足，需要 {rule.get('count')} 题，只有 {len(questions)} 题")
            
            # 随机选择指定数量的题目
            selected_questions = random.sample(questions, rule.get('count', 1))
            
            # 添加到试卷
            for question in selected_questions:
                paper_question = PaperQuestion(
                    paper_id=paper.id,
                    question_id=question.id,
                    question_order=question_order,
                    score=rule.get('score_per_question', 5.0),
                    section_name=rule.get('section_name', '')
                )
                self.db_session.add(paper_question)
                question_order += 1
        
        self.db_session.commit()
        return paper
    
    def _select_questions_by_rule(self, rule: Dict) -> List[Question]:
        """根据规则选择题目"""
        query = self.db_session.query(Question)
        
        # 按题型筛选
        if rule.get('question_type'):
            query = query.filter(Question.question_type_code == rule['question_type'])
        
        # 按难度筛选
        if rule.get('difficulty'):
            query = query.filter(Question.difficulty_code == rule['difficulty'])
        
        # 按一致性筛选（可选）
        if rule.get('consistency'):
            query = query.filter(Question.consistency_code == rule['consistency'])
        
        return query.all()
    
    def _get_default_rules(self) -> List[Dict]:
        """获取默认组题规则"""
        return [
            {
                "question_type": "B",
                "difficulty": "3",
                "count": 10,
                "score_per_question": 5.0,
                "section_name": "单选题"
            },
            {
                "question_type": "G",
                "difficulty": "3",
                "count": 5,
                "score_per_question": 8.0,
                "section_name": "多选题"
            },
            {
                "question_type": "C",
                "difficulty": "3",
                "count": 5,
                "score_per_question": 2.0,
                "section_name": "判断题"
            },
            {
                "question_type": "T",
                "difficulty": "3",
                "count": 3,
                "score_per_question": 5.0,
                "section_name": "填空题"
            },
            {
                "question_type": "D",
                "difficulty": "3",
                "count": 2,
                "score_per_question": 10.0,
                "section_name": "简答题"
            }
        ]
    
    def generate_paper_by_difficulty_distribution(self,
                                                paper_name: str,
                                                paper_description: str = "",
                                                total_score: float = 100.0,
                                                duration: int = 120,
                                                difficulty_distribution: Optional[Dict[str, float]] = None) -> Paper:
        """
        按难度分布生成试卷
        
        Args:
            paper_name: 试卷名称
            paper_description: 试卷描述
            total_score: 试卷总分
            duration: 考试时长（分钟）
            difficulty_distribution: 难度分布，格式：{"1": 0.1, "2": 0.2, "3": 0.4, "4": 0.2, "5": 0.1}
        """
        if difficulty_distribution is None:
            difficulty_distribution = {"1": 0.1, "2": 0.2, "3": 0.4, "4": 0.2, "5": 0.1}
        
        # 计算每种难度的题目数量
        total_questions = 25  # 默认总题数
        difficulty_counts = {}
        for difficulty, ratio in difficulty_distribution.items():
            difficulty_counts[difficulty] = int(total_questions * ratio)
        
        # 构建规则
        rules = []
        question_types = ["B", "G", "C", "T", "D"]  # 支持的题型
        scores = {"B": 4.0, "G": 6.0, "C": 2.0, "T": 4.0, "D": 8.0}  # 各题型分值
        
        for difficulty, count in difficulty_counts.items():
            if count > 0:
                # 为每种难度分配题型
                for i, question_type in enumerate(question_types):
                    if i < count:
                        rules.append({
                            "question_type": question_type,
                            "difficulty": difficulty,
                            "count": 1,
                            "score_per_question": scores[question_type],
                            "section_name": f"难度{difficulty}-{self._get_question_type_name(question_type)}"
                        })
        
        return self.generate_paper_by_rules(
            paper_name=paper_name,
            paper_description=paper_description,
            total_score=total_score,
            duration=duration,
            rules=rules
        )
    
    def _get_question_type_name(self, question_type_code: str) -> str:
        """获取题型名称，能处理 'B' 或 'B（单选题）' 等格式"""
        if not question_type_code:
            return "未知题型"
        
        # 提取括号前的代码，例如从 "B（单选题）" 中提取 "B"
        clean_code = question_type_code.split('（')[0].strip()

        type_names = {
            "B": "单选题",
            "G": "多选题", 
            "C": "判断题",
            "T": "填空题",
            "D": "简答题",
            "U": "计算题",
            "W": "论述题",
            "E": "案例分析",
            "F": "综合题"
        }
        return type_names.get(clean_code, "未知题型")
    
    def get_paper_statistics(self, paper_id: str) -> Dict:
        """获取试卷统计信息"""
        paper = self.db_session.query(Paper).filter(Paper.id == paper_id).first()
        if not paper:
            return {}
        
        paper_questions = self.db_session.query(PaperQuestion).filter(
            PaperQuestion.paper_id == paper_id
        ).order_by(PaperQuestion.question_order).all()
        
        # 统计信息
        stats = {
            "total_questions": len(paper_questions),
            "total_score": sum(pq.score for pq in paper_questions),
            "question_types": {},
            "difficulty_distribution": {},
            "sections": {}
        }
        
        for pq in paper_questions:
            question = pq.question
            
            # 题型统计
            question_type = question.question_type_code
            if question_type not in stats["question_types"]:
                stats["question_types"][question_type] = {
                    "count": 0,
                    "total_score": 0,
                    "name": self._get_question_type_name(question_type)
                }
            stats["question_types"][question_type]["count"] += 1
            stats["question_types"][question_type]["total_score"] += pq.score
            
            # 难度分布
            difficulty = question.difficulty_code
            if difficulty not in stats["difficulty_distribution"]:
                stats["difficulty_distribution"][difficulty] = {
                    "count": 0,
                    "total_score": 0
                }
            stats["difficulty_distribution"][difficulty]["count"] += 1
            stats["difficulty_distribution"][difficulty]["total_score"] += pq.score
            
            # 章节统计
            section_name = pq.section_name or "未分类"
            if section_name not in stats["sections"]:
                stats["sections"][section_name] = {
                    "count": 0,
                    "total_score": 0
                }
            stats["sections"][section_name]["count"] += 1
            stats["sections"][section_name]["total_score"] += pq.score
        
        return stats
    
    def export_paper_to_text(self, paper_id: str) -> str:
        """导出试卷为文本格式"""
        paper = self.db_session.query(Paper).filter(Paper.id == paper_id).first()
        if not paper:
            return "试卷不存在"
        
        paper_questions = self.db_session.query(PaperQuestion).filter(
            PaperQuestion.paper_id == paper_id
        ).order_by(PaperQuestion.question_order).all()
        
        # 生成试卷文本
        text = f"试卷名称：{paper.name}\n"
        text += f"试卷描述：{paper.description}\n"
        text += f"总分：{paper.total_score}分\n"
        text += f"考试时长：{paper.duration}分钟\n"
        text += f"难度等级：{paper.difficulty_level}\n"
        text += "=" * 50 + "\n\n"
        
        current_section = ""
        for i, pq in enumerate(paper_questions, 1):
            question = pq.question
            
            # 章节标题
            if pq.section_name and pq.section_name != current_section:
                current_section = pq.section_name
                text += f"\n{current_section}\n"
                text += "-" * 30 + "\n"
            
            # 题目内容
            text += f"{i}. ({pq.score}分) {question.stem}\n"
            
            # 选项（如果有）
            if question.option_a:
                text += f"   A. {question.option_a}\n"
            if question.option_b:
                text += f"   B. {question.option_b}\n"
            if question.option_c:
                text += f"   C. {question.option_c}\n"
            if question.option_d:
                text += f"   D. {question.option_d}\n"
            if question.option_e:
                text += f"   E. {question.option_e}\n"
            
            text += "\n"
        
        return text

    def generate_paper_by_knowledge_distribution(self, paper_name, paper_structure, knowledge_distribution, **kwargs):
        """
        核心组卷逻辑：根据题库、题型、知识点分布生成试卷
        """
        import time
        from sqlalchemy.exc import OperationalError

        # 重试机制，处理数据库锁定问题
        max_retries = 3
        retry_delay = 1

        for attempt in range(max_retries):
            try:
                paper = Paper(
                    name=paper_name,
                    description=kwargs.get('paper_description', f"基于知识点分布自动生成的试卷"),
                    total_score=kwargs.get('total_score', 100.0),
                    duration=kwargs.get('duration', 120),
                    difficulty_level=kwargs.get('difficulty_level', '中等'),
                )
                self.db_session.add(paper)
                self.db_session.flush()

                question_order = 1
                for rule in paper_structure:
                    bank_name = rule['question_bank_name']
                    q_type = rule['question_type']
                    count = rule['count']
                    score = rule['score_per_question']

                    # 核心修改：直接筛选题目，不依赖QuestionBank表
                    # 因为当前系统中所有题目都在同一个题库中
                    base_query = self.db_session.query(Question).filter(
                        Question.question_type_code == q_type
                    )

                    available_questions = base_query.all()

                    if len(available_questions) < count:
                        raise ValueError(f"题库 '{bank_name}' 中题型 '{q_type}' 的题目不足。需要 {count} 道，但只有 {len(available_questions)} 道。")

                    selected_questions = random.sample(available_questions, count)

                    for q in selected_questions:
                        pq = PaperQuestion(
                            paper_id=paper.id,
                            question_id=q.id,
                            question_order=question_order,
                            score=score,
                            section_name=f"{self._get_question_type_name(q_type)}"
                        )
                        self.db_session.add(pq)
                        question_order += 1

                self.db_session.commit()
                return paper

            except OperationalError as e:
                if "database is locked" in str(e) and attempt < max_retries - 1:
                    print(f"数据库锁定，第{attempt + 1}次重试...")
                    self.db_session.rollback()
                    time.sleep(retry_delay * (attempt + 1))
                    continue
                else:
                    self.db_session.rollback()
                    raise e
            except Exception as e:
                self.db_session.rollback()
                raise e

    def export_paper_to_docx(self, paper_id: str):
        """将试卷导出为 DOCX 格式，包含题型分组和格式化标题"""
        paper = self.db_session.query(Paper).filter(Paper.id == paper_id).first()
        if not paper:
            raise ValueError("试卷未找到")

        doc = Document()
        doc.add_heading(str(paper.name), level=1)
        
        meta_info = f"总分：{paper.total_score or 100.0}分    时长：{paper.duration or 120}分钟    难度：{paper.difficulty_level or '中等'}"
        doc.add_paragraph(meta_info)
        doc.add_paragraph("--------------------------------------------------")

        paper_questions = self.db_session.query(PaperQuestion).filter(PaperQuestion.paper_id == paper.id).order_by(PaperQuestion.question_order).all()

        if not paper_questions:
            doc.add_paragraph("该试卷内没有题目。")
        else:
            # 按题型分组
            sections = {}
            for pq in paper_questions:
                q = pq.question
                q_type_code = q.question_type_code or "UN"
                q_type_name = self._get_question_type_name(q_type_code)
                
                if q_type_name not in sections:
                    sections[q_type_name] = {"questions": [], "total_score": 0.0}
                
                sections[q_type_name]["questions"].append(pq)
                sections[q_type_name]["total_score"] += pq.score

            # 写入每个题型分组
            section_counter = 1
            roman_numerals = {1: '一', 2: '二', 3: '三', 4: '四', 5: '五', 6: '六'}
            
            for section_name, data in sections.items():
                questions_in_section = data["questions"]
                count = len(questions_in_section)
                score_per_question = questions_in_section[0].score if count > 0 else 0
                section_total_score = data["total_score"]
                
                section_title = f"{roman_numerals.get(section_counter, section_counter)}、{section_name}（共{count}题，每题{score_per_question}分，计{section_total_score}分）"
                doc.add_heading(section_title, level=2)

                for pq in questions_in_section:
                    q = pq.question
                    p = doc.add_paragraph()
                    p.add_run(f"第{pq.question_order}题. ").bold = True
                    stem_text = q.stem if q.stem is not None else ""
                    p.add_run(stem_text)

                    # 添加选项，并为单选题排除 E 选项
                    is_single_choice = q.question_type_code and q.question_type_code.startswith('B')
                    
                    if q.option_a is not None: p.add_run(f"\nA. {q.option_a}")
                    if q.option_b is not None: p.add_run(f"\nB. {q.option_b}")
                    if q.option_c is not None: p.add_run(f"\nC. {q.option_c}")
                    if q.option_d is not None: p.add_run(f"\nD. {q.option_d}")
                    if not is_single_choice and q.option_e is not None:
                        p.add_run(f"\nE. {q.option_e}")
                    
                    doc.add_paragraph() # 添加空行
                section_counter += 1

        output = BytesIO()
        doc.save(output)
        output.seek(0)
        return output 